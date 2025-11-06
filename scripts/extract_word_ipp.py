#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script pour extraire les lésions et taux IPP depuis le fichier Word IPP.docx
"""

import json
import re
import os

try:
    from docx import Document
except ImportError:
    print("❌ Module python-docx non installé.")
    print("Installation: pip install python-docx")
    exit(1)

def extract_lesions_from_word(docx_path):
    """Extrait les lésions depuis un fichier Word"""
    
    print(f"📖 Lecture du fichier Word: {docx_path}")
    
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"❌ Erreur lors de l'ouverture du fichier: {e}")
        return []
    
    lesions = []
    total_paragraphs = len(doc.paragraphs)
    print(f"📄 Total de paragraphes: {total_paragraphs}")
    
    # Patterns pour détecter les lésions et taux
    # Format attendu: "Nom de la lésion: 25%" ou "Nom: 10 à 30" ou "Nom [10-30]"
    patterns = [
        r'([A-ZÀÂÄÉÈÊËÏÎÔÙÛÜÇ][a-zàâäéèêëïîôùûüç\s\-\']+(?:[a-zàâäéèêëïîôùûüç\s\-\']+)*)\s*[:\.]?\s*(\d{1,3})\s*%',  # "Lésion: 25%"
        r'([A-ZÀÂÄÉÈÊËÏÎÔÙÛÜÇ][a-zàâäéèêëïîôùûüç\s\-\']+(?:[a-zàâäéèêëïîôùûüç\s\-\']+)*)\s*[:\.]?\s*(\d{1,3})\s*à\s*(\d{1,3})',  # "Lésion: 10 à 30"
        r'([A-ZÀÂÄÉÈÊËÏÎÔÙÛÜÇ][a-zàâäéèêëïîôùûüç\s\-\']+(?:[a-zàâäéèêëïîôùûüç\s\-\']+)*)\s*\[(\d{1,3})\s*-\s*(\d{1,3})\]',  # "Lésion [10-30]"
    ]
    
    full_text = ""
    
    # Extraire tout le texte
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text += text + "\n"
    
    # Sauvegarder le texte complet
    output_dir = "extracted_bareme"
    os.makedirs(output_dir, exist_ok=True)
    
    text_output = os.path.join(output_dir, "ipp_word_full_text.txt")
    with open(text_output, 'w', encoding='utf-8') as f:
        f.write(full_text)
    print(f"📄 Texte complet sauvegardé: {text_output}")
    print(f"📊 Caractères extraits: {len(full_text)}")
    
    # Extraire également depuis les tableaux
    table_text = ""
    if doc.tables:
        print(f"📊 Tables trouvées: {len(doc.tables)}")
        for table_idx, table in enumerate(doc.tables):
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text:
                    table_text += row_text + "\n"
                    full_text += row_text + "\n"
    
    if table_text:
        table_output = os.path.join(output_dir, "ipp_word_tables.txt")
        with open(table_output, 'w', encoding='utf-8') as f:
            f.write(table_text)
        print(f"📊 Tableaux sauvegardés: {table_output}")
    
    # Analyser le texte avec les patterns
    print(f"\n🔍 Recherche de lésions dans le texte...")
    
    for line in full_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        
        for pattern in patterns:
            matches = re.finditer(pattern, line)
            for match in matches:
                groups = match.groups()
                lesion_name = groups[0].strip()
                
                # Éviter les faux positifs (titres, headers)
                if len(lesion_name) < 5 or len(lesion_name) > 150:
                    continue
                
                if len(groups) == 2:  # Taux fixe
                    rate = int(groups[1])
                    lesion = {
                        'name': lesion_name,
                        'rate': rate,
                        'type': 'fixed',
                        'source_line': line
                    }
                    lesions.append(lesion)
                    
                elif len(groups) == 3:  # Fourchette
                    min_rate = int(groups[1])
                    max_rate = int(groups[2])
                    
                    # Vérifier que la fourchette est valide
                    if min_rate < max_rate:
                        lesion = {
                            'name': lesion_name,
                            'min_rate': min_rate,
                            'max_rate': max_rate,
                            'type': 'range',
                            'source_line': line
                        }
                        lesions.append(lesion)
    
    # Dédupliquer les lésions
    unique_lesions = []
    seen_names = set()
    
    for lesion in lesions:
        name_key = lesion['name'].lower().strip()
        if name_key not in seen_names:
            seen_names.add(name_key)
            unique_lesions.append(lesion)
    
    print(f"\n🔍 Lésions détectées: {len(unique_lesions)}")
    
    # Sauvegarder les lésions
    json_output = os.path.join(output_dir, "ipp_word_lesions.json")
    with open(json_output, 'w', encoding='utf-8') as f:
        json.dump(unique_lesions, f, ensure_ascii=False, indent=2)
    print(f"💾 Lésions sauvegardées: {json_output}")
    
    # Afficher un échantillon
    if unique_lesions:
        print(f"\n📋 Échantillon des premières lésions:")
        for i, lesion in enumerate(unique_lesions[:10], 1):
            if lesion['type'] == 'fixed':
                print(f"  {i}. {lesion['name']}: {lesion['rate']}%")
            else:
                print(f"  {i}. {lesion['name']}: {lesion['min_rate']}-{lesion['max_rate']}%")
    
    return unique_lesions

if __name__ == "__main__":
    docx_path = r"c:\Users\HICHAME\Desktop\Guide du médecin conseil\IPP.docx"
    
    if not os.path.exists(docx_path):
        print(f"❌ Fichier non trouvé: {docx_path}")
        exit(1)
    
    lesions = extract_lesions_from_word(docx_path)
    
    print(f"\n✅ Extraction terminée!")
    print(f"📊 Total lésions extraites: {len(lesions)}")
