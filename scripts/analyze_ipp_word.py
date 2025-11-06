#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Analyse médico-légale du fichier IPP.docx
Examine le contenu avant extraction pour validation
"""

try:
    from docx import Document
except ImportError:
    print("❌ Module python-docx requis: pip install python-docx")
    exit(1)

docx_path = r"c:\Users\HICHAME\Desktop\Guide du médecin conseil\IPP.docx"

print("📖 ANALYSE MÉDICO-LÉGALE DU FICHIER IPP.docx")
print("=" * 80)

try:
    doc = Document(docx_path)
    
    print(f"\n📊 STRUCTURE DU DOCUMENT:")
    print(f"  - Paragraphes: {len(doc.paragraphs)}")
    print(f"  - Tables: {len(doc.tables)}")
    
    # Compter le texte total
    total_chars = sum(len(p.text) for p in doc.paragraphs)
    print(f"  - Caractères totaux: {total_chars:,}")
    
    # Afficher les premiers paragraphes
    print(f"\n📄 PREMIERS 30 PARAGRAPHES (aperçu du contenu):")
    print("-" * 80)
    
    count = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and count < 30:
            # Limiter la longueur d'affichage
            display_text = text[:150] + "..." if len(text) > 150 else text
            print(f"{count+1}. {display_text}")
            count += 1
    
    # Analyser les tables si présentes
    if doc.tables:
        print(f"\n📊 ANALYSE DES TABLES:")
        print("-" * 80)
        for table_idx, table in enumerate(doc.tables[:3], 1):  # Premières 3 tables
            print(f"\n  Table {table_idx}: {len(table.rows)} lignes × {len(table.columns)} colonnes")
            print(f"  Échantillon des 5 premières lignes:")
            for row_idx, row in enumerate(table.rows[:5], 1):
                row_text = " | ".join([cell.text.strip()[:30] for cell in row.cells])
                print(f"    {row_idx}. {row_text}")
    
    # Rechercher des patterns médicaux communs
    print(f"\n🔍 RECHERCHE DE PATTERNS MÉDICO-LÉGAUX:")
    print("-" * 80)
    
    full_text = "\n".join([p.text for p in doc.paragraphs])
    
    patterns = {
        "Fracture": full_text.lower().count("fracture"),
        "Amputation": full_text.lower().count("amputation"),
        "Luxation": full_text.lower().count("luxation"),
        "Hernie": full_text.lower().count("hernie"),
        "Perte": full_text.lower().count("perte"),
        "Ankylose": full_text.lower().count("ankylose"),
        "Paralysie": full_text.lower().count("paralysie"),
        "Taux IPP / %": full_text.count("%"),
        "Fourchette (à)": full_text.count(" à "),
        "Main dominante": full_text.lower().count("main dominante"),
    }
    
    for pattern, count in patterns.items():
        if count > 0:
            print(f"  ✓ {pattern}: {count} occurrences")
    
    print(f"\n✅ Analyse terminée - Fichier prêt pour extraction")
    
except Exception as e:
    print(f"\n❌ ERREUR: {e}")
    import traceback
    traceback.print_exc()
